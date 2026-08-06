"""Render ROADMAP-TO-PRODUCTION.md to the Word document on the Desktop.

Handles the subset of Markdown the roadmap actually uses: ATX headings, tables
with a separator row, bullets, blockquotes, horizontal rules, and inline
**bold** / *italic* / `code`.
"""
from __future__ import annotations

import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

SRC = sys.argv[1]
DST = sys.argv[2]

INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+?\*)")


def add_runs(par, text: str) -> None:
    for piece in INLINE.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**") and len(piece) > 4:
            par.add_run(piece[2:-2]).bold = True
        elif piece.startswith("`") and piece.endswith("`") and len(piece) > 2:
            r = par.add_run(piece[1:-1])
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
            r.font.color.rgb = RGBColor(0xB0, 0x30, 0x60)
        elif piece.startswith("*") and piece.endswith("*") and len(piece) > 2:
            par.add_run(piece[1:-1]).italic = True
        else:
            par.add_run(piece)


def is_table_sep(line: str) -> bool:
    return bool(re.fullmatch(r"\s*\|[\s:|-]+\|\s*", line))


def cells(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


doc = Document()
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)

raw = open(SRC, encoding="utf-8").read().splitlines()

# Unwrap hard-wrapped prose FIRST.
#
# The source wraps at ~80 columns, so a bold span routinely straddles a line
# break: "**Section A is\nthe only section ...**". Rendering line-by-line splits
# the markers and leaves literal asterisks in the Word file — and produces one
# Word paragraph per source line, which reads as ragged nonsense. Consecutive
# non-blank prose lines are ONE paragraph, which is what Markdown means anyway.
SPECIAL = re.compile(r"^(#|>|\||---\s*$|\s*[-*] |\s*\d+\. )")
lines: list[str] = []
buf: list[str] = []


def flush() -> None:
    if buf:
        lines.append(" ".join(x.strip() for x in buf))
        buf.clear()


for ln in raw:
    if not ln.strip():
        flush()
        lines.append("")
    elif SPECIAL.match(ln):
        # A bullet or numbered item continues onto indented following lines.
        if buf and re.match(r"^\s{2,}\S", ln) and not SPECIAL.match(ln.strip()):
            buf.append(ln)
        # Consecutive blockquote lines are ONE quote. Flushing each separately
        # split any bold span straddling the wrap and left literal asterisks in
        # the Word file — the same defect the prose unwrapping above fixes.
        elif ln.lstrip().startswith(">") and buf and buf[0].lstrip().startswith(">"):
            buf.append(re.sub(r"^\s*>\s?", "", ln))
        else:
            flush()
            buf.append(ln)
            if ln.lstrip().startswith(("#", "|")) or ln.strip() == "---":
                flush()
    else:
        if buf and SPECIAL.match(buf[0]) and not buf[0].lstrip().startswith(("#", ">", "|")):
            buf.append(ln)          # continuation of a bullet
        else:
            buf.append(ln)
flush()

i = 0
while i < len(lines):
    line = lines[i]

    # Table: a pipe row followed by a separator row.
    if line.strip().startswith("|") and i + 1 < len(lines) and is_table_sep(lines[i + 1]):
        header = cells(line)
        i += 2
        rows = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            rows.append(cells(lines[i]))
            i += 1
        width = max([len(header)] + [len(r) for r in rows])
        t = doc.add_table(rows=1, cols=width)
        t.style = "Light Grid Accent 1"
        for n, h in enumerate(header):
            c = t.rows[0].cells[n]
            c.text = ""
            add_runs(c.paragraphs[0], h)
            for r in c.paragraphs[0].runs:
                r.bold = True
        for row in rows:
            rc = t.add_row().cells
            for n in range(width):
                rc[n].text = ""
                add_runs(rc[n].paragraphs[0], row[n] if n < len(row) else "")
        doc.add_paragraph()
        continue

    if line.startswith("#"):
        lvl = len(line) - len(line.lstrip("#"))
        doc.add_heading(line[lvl:].strip().replace("**", ""), level=min(lvl, 4))
    elif line.strip() == "---":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run("· · ·").font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    elif line.startswith("> "):
        p = doc.add_paragraph(style="Intense Quote")
        add_runs(p, line[2:])
    elif re.match(r"^\s*[-*] ", line):
        indent = (len(line) - len(line.lstrip())) // 2
        p = doc.add_paragraph(style="List Bullet")
        if indent:
            p.paragraph_format.left_indent = Pt(18 * (indent + 1))
        add_runs(p, re.sub(r"^\s*[-*] ", "", line))
    elif re.match(r"^\s*\d+\. ", line):
        p = doc.add_paragraph(style="List Number")
        add_runs(p, re.sub(r"^\s*\d+\. ", "", line))
    elif line.strip():
        add_runs(doc.add_paragraph(), line)

    i += 1

doc.save(DST)
print(f"wrote {DST}")
